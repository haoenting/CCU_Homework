from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.common.exceptions import NoSuchElementException
from selenium.webdriver.support.ui import Select
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from bs4 import BeautifulSoup
from docx import Document
import unicodedata
from docx.enum.text import WD_BREAK
import time
import re

courts = ['SCD', 'MLD', 'TCD', 'CHD',
          'NTD', 'ULD', 'CYD', 'TND', 'KSD', 'CTD', 'HLD', 'TTD', 'PTD', 'PHD', 'KMD', 'LCD']
for year in range(107, 108):
    for court in courts:
        options = webdriver.ChromeOptions()
        options.add_experimental_option('excludeSwitches', ['enable-logging'])
        driver_path = 'D:/chromedriver'  # 替换为你的 WebDriver 路径
        driver = webdriver.Chrome(executable_path=driver_path, options=options)
        url = 'https://judgment.judicial.gov.tw/FJUD/Default_AD.aspx'
        driver.get(url)

        court_select = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.ID, 'jud_court'))
        )

        court_select = Select(court_select)
        court_select.select_by_value(court)
        context = driver.find_element(By.ID, 'dy1')
        context.send_keys(year)
        context = driver.find_element(By.ID, 'dm1')
        context.send_keys('1')
        context = driver.find_element(By.ID, 'dd1')
        context.send_keys('1')
        context = driver.find_element(By.ID, 'dy2')
        context.send_keys(year)
        context = driver.find_element(By.ID, 'dm2')
        context.send_keys('12')
        context = driver.find_element(By.ID, 'dd2')
        context.send_keys('31')
        context = driver.find_element(By.ID, 'jud_kw')
        context.send_keys('公約+條約+議定書+協定-條約定-協議-合議-住戶管理公約-制式條約')

        search_button = driver.find_element(By.ID, 'btnQry')
        search_button.click()

        # 切换到iframe
        iframe = driver.find_element(By.XPATH, '//iframe[@id="iframe-data"]')
        driver.switch_to.frame(iframe)

        elements = []
        for i in range(0, 25):
            # 解析iframe内部的内容
            page_source = driver.page_source
            soup = BeautifulSoup(page_source, 'html.parser')

            # 拿到所有的標題
            elements += soup.find_all('a', class_='hlTitle_scroll')

            try:
                button = driver.find_element(By.ID, 'hlNext')
                button.click()
            except NoSuchElementException:
                print("No such button")
                break

        index = int(1)
        print('total :', len(elements))
        for element in elements:
            title = element.text
            link = "https://judgment.judicial.gov.tw/FJUD/" + element['href']
            driver.get(link)
            driver.implicitly_wait(10)
            judgment_soup = BeautifulSoup(driver.page_source, 'html.parser')

            divs = judgment_soup.find(
                'td', class_='tab_content').find_all('div')
            document = Document()
            document.add_paragraph().add_run(title).bold = True
            try:
                for div in divs:
                    if not div.text:
                        continue
                    paragraph = document.add_paragraph(div.text)
                    paragraph.add_run('\n')
            except ValueError:
                continue
            document.add_page_break()

            import os
            from datetime import datetime
            file_path = "D:/判決書"
            file_path = os.path.join(file_path, str(year))

            if not os.path.exists(file_path):
                os.makedirs(file_path)

            file_name = f'{file_path}/{title}.docx'
            document.save(file_name)

            print(index, title, "done")
            index += 1
            time.sleep(1)
        driver.switch_to.default_content()
        driver.quit()
